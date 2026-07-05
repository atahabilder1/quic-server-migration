#!/usr/bin/env python3
"""Generate DETAILED Research Task Document (Word + PDF) for QUIC Server Migration.

Matches the 5-task structure of RESEARCH_TASKS_SHORT.docx with full detail.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from datetime import date

OUTPUT_DOCX = "ideas/RESEARCH_TASKS_DETAILED.docx"

DARK_BLUE = RGBColor(0x1a, 0x36, 0x5d)
MED_BLUE = RGBColor(0x2b, 0x6c, 0xb0)
DARK_GRAY = RGBColor(0x2d, 0x37, 0x48)
GREEN = RGBColor(0x27, 0x67, 0x49)
ORANGE = RGBColor(0xc0, 0x56, 0x21)
RED = RGBColor(0xe5, 0x3e, 0x3e)


def set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    elem = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    shading.append(elem)


def add_h1(doc, text):
    h = doc.add_heading(text, level=1)
    for run in h.runs:
        run.font.color.rgb = DARK_BLUE
    return h


def add_h2(doc, text):
    h = doc.add_heading(text, level=2)
    for run in h.runs:
        run.font.color.rgb = MED_BLUE
    return h


def add_h3(doc, text):
    h = doc.add_heading(text, level=3)
    for run in h.runs:
        run.font.color.rgb = DARK_GRAY
    return h


def add_goal_box(doc, text, shade='DCE6F1'):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    p = cell.paragraphs[0]
    run_b = p.add_run("Goal: ")
    run_b.bold = True
    run_b.font.size = Pt(10)
    run_b.font.color.rgb = DARK_BLUE
    run_t = p.add_run(text)
    run_t.font.size = Pt(10)
    run_t.font.color.rgb = DARK_BLUE
    set_cell_shading(cell, shade)
    doc.add_paragraph()


def add_alert_box(doc, text, shade='FFF5F5'):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    p = cell.paragraphs[0]
    run_b = p.add_run("Key Issue: ")
    run_b.bold = True
    run_b.font.size = Pt(10)
    run_b.font.color.rgb = RED
    run_t = p.add_run(text)
    run_t.font.size = Pt(10)
    run_t.font.color.rgb = DARK_GRAY
    set_cell_shading(cell, shade)
    doc.add_paragraph()


def add_body(doc, text):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        run_b.font.size = Pt(10)
        p.add_run(text).font.size = Pt(10)
    else:
        p.add_run(text).font.size = Pt(10)
    return p


def add_table(doc, data):
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.style = 'Medium Shading 1 Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row_data in enumerate(data):
        for j, cell_text in enumerate(row_data):
            table.cell(i, j).text = cell_text
            for p in table.cell(i, j).paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()
    return table


def add_rq(doc, num, text):
    p = doc.add_paragraph(style='List Bullet')
    run_b = p.add_run(f"RQ{num}: ")
    run_b.bold = True
    run_b.font.size = Pt(10)
    p.add_run(text).font.size = Pt(10)


def add_separator(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run('\u2500' * 80)
    run.font.size = Pt(6)
    run.font.color.rgb = MED_BLUE


def build_doc():
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    style.font.color.rgb = DARK_GRAY

    # ── TITLE ──
    title = doc.add_heading('Research Tasks & Exploration Plan', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = DARK_BLUE

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run('QUIC Server-Side Migration Project')
    run.font.size = Pt(12)
    run.font.color.rgb = MED_BLUE

    dt = doc.add_paragraph()
    dt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = dt.add_run(date.today().strftime('%B %d, %Y'))
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x4a, 0x55, 0x68)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Detailed Version \u2014 Comprehensive descriptions, research questions, '
                     'comparison tables, and sub-task breakdowns.')
    run.font.size = Pt(10)
    run.italic = True
    run.font.color.rgb = MED_BLUE

    doc.add_paragraph('_' * 80)

    # ── CURRENT STATUS ──
    add_h1(doc, 'Current Status')
    add_body(doc,
        'We have a fully working cross-machine QUIC server migration implementation '
        'using modified Mozilla Neqo (Rust). Verified with Firefox over HTTP/3 on 4 physical '
        'machines (same LAN). Anik: 5 state transfer backends. '
        'Fatima: Docker/namespace setup with /tmp migration.')


    # ── RESEARCH FOCUS ──
    add_h1(doc, 'Research Focus (Per Professor)')
    add_bullet(doc,
        'Explore primary as proxy during handshake. Compare against baselines.',
        bold_prefix='QUIC Proxy (Pass-Through Handshake) \u2014 ')
    add_bullet(doc,
        'What are the real applications of QUIC migration? Server architecture, '
        'engineering perspective.',
        bold_prefix='Precise Definition of Applications \u2014 ')

    # ── OVERVIEW TABLE ──
    add_h1(doc, 'All Tasks at a Glance')
    add_table(doc, [
        ['#', 'Task', 'Type'],
        ['1', 'Docker Migration + 4 Backend Comparison\n(/tmp, TCP, Redis, QUIC)',
         'Engineering +\nBenchmarking'],
        ['2', 'QUIC Proxy (Pass-Through Handshake)\nExploration + Decryption Problem',
         'Protocol\nExploration'],
        ['3', 'Applications of QUIC Migration\n(Use Cases, Server Types, Stream LB,\nServer Chaining)',
         'Research +\nExploration'],
        ['4', 'Neqo Server & Server Architecture\n(Neqo Capabilities, nginx Comparison,\nArchitecture Requirements)',
         'Investigation +\nAnalysis'],
        ['5', 'Attack Surface vs. Benefits\nTradeoff Analysis',
         'Security\nAnalysis'],
    ])

    # ══════════════════════════════════════════════════════════════════
    # TASK 1: Docker + Backends
    # ══════════════════════════════════════════════════════════════════
    doc.add_page_break()
    add_h1(doc, 'Task 1: Docker-Based Migration with 4 Backend Comparison')
    add_goal_box(doc,
        'Convert Fatima\u2019s namespace-based implementation to Docker containers. '
        'Implement and benchmark 4 state transfer backends: /tmp (shared volume), TCP Push, '
        'Redis KV, and QUIC-based transfer. Measure latency, reliability, and migration success rate.')

    add_h2(doc, 'Background')
    add_body(doc,
        'Fatima has completed a working migration using network namespaces with /tmp (shared '
        'filesystem) as the state transfer mechanism. The migration state is ~445 bytes '
        'containing TLS secrets, connection IDs, packet numbers, and client address. The next '
        'step is to convert this to Docker containers and implement additional backends.')

    add_h2(doc, 'Sub-Tasks')
    add_bullet(doc,
        'Each server (primary, preferred) in its own container. '
        'Shared Docker network for QUIC traffic. Redis in a third container.',
        bold_prefix='1a. Convert Namespace to Docker: ')
    add_bullet(doc,
        'Primary writes state to shared Docker volume. Preferred reads from same volume. '
        'Baseline \u2014 zero network overhead. Leaves disk forensic artifacts.',
        bold_prefix='1b. /tmp Backend (Shared Volume): ')
    add_bullet(doc,
        'Primary opens TCP to preferred, pushes 445 bytes. Already implemented on physical '
        'machines \u2014 port to Docker. Observable side-channel (TCP connection visible). '
        'Lowest latency (~1ms).',
        bold_prefix='1c. TCP Push Backend: ')
    add_bullet(doc,
        'State in Redis (separate container). Primary SETs, preferred GETs. '
        'Decoupled \u2014 supports lazy retrieval. Best multi-server scalability.',
        bold_prefix='1d. Redis KV Backend: ')
    add_bullet(doc,
        'QUIC connection between primary and preferred for state transfer. '
        'Encrypted, indistinguishable from regular QUIC traffic. '
        'Hardest for IDS to detect. New implementation required.',
        bold_prefix='1e. QUIC-Based State Transfer (NEW): ')

    add_h2(doc, 'Comparison Metrics')
    for m in [
        'Transfer latency (state export to preferred ready)',
        'Migration success rate (over 100 consecutive migrations)',
        'Time-to-first-byte after migration (client perspective)',
        'Network fingerprint visibility (tcpdump + Wireshark analysis)',
        'Disk/memory forensic artifacts',
        'Scalability (N preferred servers)',
        'Failure resilience (preferred not ready scenario)',
    ]:
        add_bullet(doc, m)

    add_h2(doc, 'Expected Comparison Matrix')
    add_table(doc, [
        ['Backend', 'Latency', 'Network\nVisible?', 'Forensic\nArtifacts', 'Scalability', 'Dependencies'],
        ['/tmp\n(shared vol)', 'Lowest\n(disk I/O)', 'No', 'Yes\n(file on disk)', 'Single host\nonly', 'Shared\nstorage'],
        ['TCP Push', '~1ms\n(LAN)', 'Yes\n(TCP conn)', 'No\n(in-memory)', 'Point-to-\npoint', 'None'],
        ['Redis KV', '~2-3ms\n(LAN)', 'Yes\n(Redis proto)', 'Yes\n(Redis log)', 'Multi-\nserver', 'Redis\ninstance'],
        ['QUIC\nTransfer', '~2-5ms\n(LAN)', 'Encrypted\n(looks normal)', 'No\n(in-memory)', 'Point-to-\npoint', 'QUIC stack\non both'],
    ])

    add_h2(doc, 'Research Questions')
    for i, q in enumerate([
        'Which backend achieves the lowest migration latency?',
        'Does the choice of backend affect migration success rate?',
        'Can a network observer distinguish QUIC-based state transfer from regular QUIC traffic?',
        'Which backend is most resilient when the preferred server is not immediately ready?',
        'What is the operational overhead of each backend in a real deployment?',
    ], 1):
        add_rq(doc, i, q)

    # ══════════════════════════════════════════════════════════════════
    # TASK 2: QUIC Proxy
    # ══════════════════════════════════════════════════════════════════
    doc.add_page_break()
    add_h1(doc, 'Task 2: QUIC Proxy \u2014 Pass-Through Handshake Exploration')
    add_goal_box(doc,
        'Explore whether the primary server can act as a proxy during the QUIC/TLS 1.3 '
        'handshake, forwarding handshake messages to the preferred server in real-time so it '
        'derives its own crypto keys. Identify the decryption problem. If a practical solution '
        'exists, compare latency against Task 1 baselines.')

    add_h2(doc, 'Motivation')
    add_body(doc,
        'In our current implementation, the primary completes the full TLS 1.3 handshake, '
        'then exports ~445 bytes of crypto state to the preferred server. This creates a window '
        'where the preferred has no crypto context. A pass-through model would let the preferred '
        'participate from the beginning, potentially eliminating the state transfer delay.')

    add_h2(doc, 'How It Would Work (Ideal Case)')
    for i, s in enumerate([
        'Client sends Initial packet to the primary server.',
        'Primary handles handshake normally BUT also forwards handshake packets to preferred in real-time.',
        'Preferred silently processes the same handshake messages and derives identical TLS keys.',
        'Handshake completes. Primary advertises preferred_address.',
        'Client sends PATH_CHALLENGE to preferred.',
        'Preferred already has the keys \u2014 responds immediately. Zero delay.',
    ], 1):
        add_bullet(doc, s, bold_prefix=f'Step {i}: ')

    add_h2(doc, 'The Decryption Problem (Critical Challenge)')
    add_alert_box(doc,
        'Pure pass-through does NOT work due to TLS 1.3 ECDHE. The preferred server cannot '
        'derive traffic keys just by observing forwarded handshake packets. It lacks the '
        'primary\u2019s ephemeral private key needed for the Diffie-Hellman computation.')

    add_body(doc, 'In TLS 1.3, the handshake uses ECDHE:')
    add_bullet(doc, 'Client sends ephemeral PUBLIC key in ClientHello.')
    add_bullet(doc, 'Primary generates ephemeral key pair, sends PUBLIC key in ServerHello.')
    add_bullet(doc, 'Both compute: shared_secret = their_private \u00d7 other_public.')

    add_body(doc,
        'If primary forwards these packets, preferred sees both PUBLIC keys but does NOT have '
        'the primary\u2019s ephemeral PRIVATE key. Without it, it CANNOT compute the shared secret.')
    add_body(doc,
        'The primary would still need to send either the ephemeral private key OR the derived '
        'shared secret. This brings us back to crypto state export \u2014 just during the '
        'handshake instead of after. The latency benefit needs experimental verification.')

    add_h2(doc, 'Current vs. Pass-Through Comparison')
    add_table(doc, [
        ['Aspect', 'Current\n(Post-Handshake Export)', 'QUIC Proxy\n(Pass-Through)'],
        ['When preferred gets keys', 'After handshake', 'During handshake\n(if feasible)'],
        ['What is transferred', '~445 bytes\n(secrets, CIDs, pkt#)', 'Ephemeral key or\nshared secret'],
        ['Transfer delay', 'After handshake\n+ transfer time', 'Overlaps with\nhandshake'],
        ['Decrypt issue?', 'No \u2014 full state\nexported', 'Yes \u2014 needs\ninvestigation'],
        ['Complexity', 'Moderate', 'Higher\n(real-time fwd)'],
    ])

    add_h2(doc, 'Research Questions')
    for i, q in enumerate([
        'Can we forward enough TLS state during the handshake for preferred to derive keys?',
        'Does forwarding the ephemeral key during handshake reduce latency vs. post-handshake export?',
        'What are the security implications? Does sharing ephemeral key weaken forward secrecy?',
        'Is there a way to involve preferred in key generation from the start (coordinated ECDHE)?',
        'Compatible with TLS 1.3 session resumption and 0-RTT?',
        'What is the minimum state that must be forwarded during handshake vs. after?',
    ], 1):
        add_rq(doc, i, q)

    add_h2(doc, 'Deliverable')
    add_body(doc,
        'A feasibility report: (1) whether pass-through is possible given TLS 1.3, '
        '(2) decryption limitations, (3) if partial solution exists, '
        '(4) if feasible, latency benchmarks against Task 1 baselines.')

    # ══════════════════════════════════════════════════════════════════
    # TASK 3: Applications of QUIC Migration
    # ══════════════════════════════════════════════════════════════════
    doc.add_page_break()
    add_h1(doc, 'Task 3: Applications of QUIC Migration')
    add_goal_box(doc,
        'Explore and research the real-world applications of QUIC server-side migration. '
        'Investigate use cases and content models, server type applicability, stream-level '
        'load balancing, and server-to-server chaining. Do basic research on each area and '
        'list what should be investigated further.',
        shade='F0FFF4')

    add_body(doc,
        'What is QUIC server migration actually useful for, and where does it fall short?')

    # ── 3a ──
    add_separator(doc)
    add_h2(doc, '3a. Use Cases and Content Model')
    add_body(doc,
        'What are the concrete applications of QUIC server-side migration? For each, specify: '
        'server architecture, content model (same vs. different), migration workflow, and what '
        'additional state (beyond TLS) needs to migrate.')

    add_h3(doc, 'Same Content vs. Different Content')
    add_table(doc, [
        ['Scenario', 'Content Model', 'How It Works', 'Example'],
        ['Horizontal LB\n(Replicas)', 'Same content\non all servers', 'Client gets identical\nresponse from any', 'CDN edges,\nstateless APIs'],
        ['Vertical\nSeparation', 'Different roles,\ndifferent content', 'Primary = gateway\nPreferred = backend', 'TLS terminator\n+ app server'],
        ['Stateful\nServices', 'Same content,\nshared state', 'App state in\nexternal store', 'Shopping cart,\nsessions'],
        ['Malicious', 'Different\n(attacker ctrl)', 'Preferred serves\nattacker content', 'QUIC-Exfil'],
    ])

    add_h3(doc, 'Application: Seamless Load Balancing (Horizontal)')
    add_body(doc,
        'Both servers are identical replicas. Primary handles TLS handshake, migrates to '
        'a less-loaded replica. After migration, replica communicates directly with client '
        '\u2014 no proxy in the data path. Key advantage over nginx: no proxy bottleneck.')

    add_h3(doc, 'Application: Vertical Separation')
    add_body(doc,
        'Primary is a lightweight TLS gateway. It handles handshake/auth, then migrates to '
        'a heavyweight application backend. Different roles, potentially different software. '
        'This is the \u2018vertical\u2019 split the professor mentioned.')

    add_h3(doc, 'Application: Zero-Downtime Maintenance')
    add_body(doc,
        'Active connections migrated to standby before primary goes offline. '
        'No client errors. Supports rolling updates.')

    add_h3(doc, 'Application: Geographic Optimization / Failover')
    add_body(doc,
        'Mid-session migration to better server. Sub-second failover without TCP reconnect. '
        'Not possible with DNS after connection established.')

    add_h3(doc, 'What State Needs to Migrate?')
    add_table(doc, [
        ['State Type', 'Currently\nMigrated?', 'Size', 'Needed For'],
        ['TLS secrets', 'Yes', '~128 bytes', 'All scenarios'],
        ['Connection IDs', 'Yes', '~40 bytes', 'All scenarios'],
        ['Packet numbers', 'Yes', '~16 bytes', 'All scenarios'],
        ['HTTP/3 stream state', 'No', 'Variable', 'Mid-request migration'],
        ['Application session', 'No', 'Variable', 'Stateful services'],
        ['Request context\n(URL, headers)', 'No', 'Variable', 'Vertical separation'],
        ['Flow control state', 'No', '~32 bytes', 'High-throughput'],
    ])

    add_h3(doc, 'What to Check Further')
    for q in [
        'For each application, what is the minimum state that must migrate?',
        'Can we define a generic \u2018application state export\u2019 API for different server types?',
        'Which application has the strongest case for a paper contribution?',
        'Are there real-world deployments we can reference or compare against?',
    ]:
        add_bullet(doc, q)

    # ── 3b ──
    add_separator(doc)
    add_h2(doc, '3b. Server Type Applicability')
    add_body(doc,
        'Is QUIC migration equally applicable to all server types, or does each need '
        'separate implementation? Our 445-byte migration = TLS state only. '
        'What about stateful servers?')

    add_table(doc, [
        ['Server Type', 'Stateless?', 'App State\nto Migrate', 'Migration\nComplexity', 'Example'],
        ['Static web', 'Yes', 'None', 'Easy \u2014 current\nimpl sufficient', 'nginx HTML'],
        ['REST API', 'Mostly', 'Session (maybe)', 'Easy if\nstateless', 'Microservices'],
        ['Streaming\n(Netflix)', 'No', 'Stream position,\nbuffer, bitrate', 'Hard \u2014 need\nplayback state', 'Video, live'],
        ['Database-\nbacked', 'No', 'DB conn,\ntransaction', 'Hard \u2014 DB\ndon\u2019t migrate', 'E-commerce'],
        ['WebSocket /\nreal-time', 'No', 'Session,\nsubscriptions', 'Hard \u2014\ncomplex state', 'Chat, gaming'],
        ['CDN edge', 'Yes (cache)', 'Cache (optional)', 'Easy if cached\non both', 'Cloudflare'],
    ])

    add_h3(doc, 'Key Insight')
    add_body(doc,
        'For stateless servers (static web, REST APIs, CDN), our current TLS-only migration '
        'is sufficient. For stateful servers (streaming, database, real-time), additional '
        'application state must transfer \u2014 our implementation doesn\u2019t handle this.')

    add_h3(doc, 'What to Check Further')
    for q in [
        'Can we define a generic \u2018application state export\u2019 interface for different server types?',
        'Is there a clean separation: transport migration (QUIC) vs. app migration (HTTP/3)?',
        'One framework or separate implementations per type?',
        'For Netflix-style streaming: what exactly needs to transfer? (bitrate, position, buffer)',
        'For database: can we migrate the QUIC connection but keep the DB connection separate?',
    ]:
        add_bullet(doc, q)

    # ── 3c ──
    add_separator(doc)
    add_h2(doc, '3c. Stream-Level Load Balancing \u2014 Why It Doesn\u2019t Work')
    add_body(doc,
        'QUIC has independent multiplexed streams, which raises a natural question: could '
        'individual streams be routed to different servers? We analyze why this is not feasible '
        'with the current protocol.')

    add_h3(doc, 'Why Per-Stream Routing Breaks Down')
    add_body(doc,
        'Despite QUIC streams being logically independent, they share critical '
        'connection-level state that cannot be split:')
    add_bullet(doc,
        'All streams use the same TLS keys. A single QUIC packet can carry frames from '
        'multiple streams \u2014 you cannot decrypt one stream without decrypting the entire packet.',
        bold_prefix='Shared encryption: ')
    add_bullet(doc,
        'Packet numbers are per-connection, not per-stream. Multiple servers would need to '
        'coordinate packet number allocation to avoid collisions.',
        bold_prefix='Shared packet numbers: ')
    add_bullet(doc,
        'QUIC runs one congestion controller per connection. Splitting streams across servers '
        'breaks congestion feedback.',
        bold_prefix='Shared congestion control: ')
    add_bullet(doc,
        'ACKs are per-connection. Server A cannot acknowledge packets that Server B received.',
        bold_prefix='Shared ACK state: ')
    add_bullet(doc,
        'RFC 9000 preferred_address migrates the entire connection, not individual streams. '
        'No per-stream mechanism exists.')

    add_h3(doc, 'Comparison: Split Streams vs. Separate Connections')
    add_table(doc, [
        ['Approach', 'Feasibility', 'Trade-off'],
        ['Per-stream routing\n(one connection)', 'Not feasible\nwith current QUIC', 'Would need new protocol,\nbreaks crypto/congestion'],
        ['Separate QUIC\nconnections per service', 'Works today', 'More handshakes,\nbut each connection\nis self-contained'],
    ])

    add_h3(doc, 'What to Check Further')
    for q in [
        'Is there existing research on per-stream routing? How do they handle shared state?',
        'Could a proxy model work where one server forwards specific stream frames to backends?',
        'Is the overhead of multiple separate connections actually a problem in practice?',
    ]:
        add_bullet(doc, q)

    # ── 3d ──
    add_separator(doc)
    add_h2(doc, '3d. Server-to-Server Migration in Multi-Tier Architectures')
    add_body(doc,
        'In multi-tier architectures, a front-end server receives a client request and makes '
        'its own QUIC connection to a backend. If the backend advertises preferred_address, '
        'the front-end (acting as a QUIC client) follows the migration. This is a single-hop '
        'migration on the backend side, not a chain.')

    add_body(doc,
        'Client \u2192 Server A (front-end) \u2192 Server B (backend) \u2192 Server C (migrated backend)\n'
        'The client\u2019s connection to A does not migrate. A\u2019s connection to B migrates to C. '
        'These are two independent connections, each with at most one migration.')

    add_h3(doc, 'Limitations')
    add_bullet(doc,
        'preferred_address is a one-time handshake parameter. Once a connection migrates, '
        'it cannot migrate again \u2014 there is no mechanism for the preferred server to '
        'advertise a second preferred address.')
    add_bullet(doc,
        'Cascading failover (B fails, so migrate to C) requires B to proactively migrate '
        'before failure. A dead server cannot export state. This is the same as '
        'zero-downtime maintenance (covered in 3a).')
    add_bullet(doc,
        'Multi-hop relay (US \u2192 EU \u2192 Asia) would require repeated migrations on the '
        'same connection, which the protocol does not support.')

    add_h3(doc, 'What Is Feasible')
    add_bullet(doc,
        'Client\u2192A can migrate (client-facing tier). A\u2192B can separately migrate '
        '(backend tier). Each tier manages its own connections.',
        bold_prefix='Independent per-tier migration: ')
    add_bullet(doc,
        'Front-end\u2019s connection to backend B migrates to C. Front-end follows transparently. '
        'Useful for microservice rebalancing.',
        bold_prefix='Backend rebalancing: ')

    add_h3(doc, 'What to Check Further')
    for q in [
        'Is independent per-tier migration practical? What is the cumulative latency?',
        'Does backend migration add meaningful attack surface beyond single-hop?',
        'Are there real microservice architectures where backend QUIC migration helps?',
    ]:
        add_bullet(doc, q)

    # ── Task 3 Summary ──
    add_separator(doc)
    add_h2(doc, 'Task 3 Summary: What to Explore')
    add_table(doc, [
        ['Sub-Task', 'Core Question', 'Check Further'],
        ['3a. Use Cases &\nContent Model', 'What are the real uses?\nSame vs. different content?', 'Min state per app,\ngeneric API, paper angle'],
        ['3b. Server Types', 'Works for all servers\nor just stateless?', 'App state interface,\nstreaming/DB requirements'],
        ['3c. Stream-Level LB', 'Why can\u2019t streams go\nto different servers?', 'Protocol constraints,\nseparate connections\nas alternative'],
        ['3d. Multi-Tier\nMigration', 'Can backend connections\nmigrate independently?', 'Per-tier feasibility,\nmicroservice use cases'],
    ])

    # ══════════════════════════════════════════════════════════════════
    # TASK 4: Neqo Server & Server Architecture
    # ══════════════════════════════════════════════════════════════════
    doc.add_page_break()
    add_h1(doc, 'Task 4: Neqo Server & Server Architecture')
    add_goal_box(doc,
        'Investigate what Neqo\u2019s server actually supports and why someone said '
        '\u2018Neqo doesn\u2019t support server.\u2019 Compare QUIC migration architecture against '
        'traditional nginx reverse proxy. Understand what server architecture is needed '
        'to support migration and whether it can coexist with existing infrastructure.',
        shade='FFFAF0')

    # ── 4a ──
    add_h2(doc, '4a. Neqo Server: What Does It Actually Support?')
    add_body(doc,
        'Someone said \u2018Neqo doesn\u2019t support server.\u2019 Neqo is built for Firefox (a client). '
        'The server component exists but may be limited to testing. We need to understand '
        'exactly what it can and cannot do, and whether it matters for our research.')

    add_h3(doc, 'What Neqo Server Has')
    for p in [
        'neqo-bin/src/server/ \u2014 working QUIC server (HTTP/0.9 and HTTP/3 modes)',
        'neqo-transport with Role::Server \u2014 full QUIC transport layer',
        'preferred_address support (our modification)',
        'Http3Server in neqo-http3',
        'TLS certificate configuration via CLI',
    ]:
        add_bullet(doc, p)

    add_h3(doc, 'What May Be Missing')
    for p in [
        'Production concurrency (single-threaded event loop vs. multi-threaded workers?)',
        'Full HTTP/3 features (routing, virtual hosts, content negotiation?)',
        'Performance under load (max concurrent connections?)',
        'Graceful shutdown / connection draining',
        'Integration with existing infrastructure (nginx, health checks, monitoring)',
    ]:
        add_bullet(doc, p)

    add_table(doc, [
        ['Feature', 'Neqo Server', 'Production Server\n(nginx/quiche)', 'Impact on Research'],
        ['QUIC transport', 'Full', 'Full', 'None'],
        ['HTTP/3', 'Basic (test)', 'Full production', 'Limits app testing'],
        ['preferred_address', 'Supported\n(our mod)', 'NOT implemented\nin production servers', 'This IS our\ncontribution'],
        ['Concurrency', 'Limited?', 'Thousands', 'Affects benchmarks'],
        ['Content serving', 'Static/test', 'Full web server', 'Limits app testing'],
        ['nginx integration', 'Unknown', 'Native', 'Needs investigation'],
    ])

    add_h3(doc, 'What to Check Further')
    for q in [
        'Read neqo-bin/src/server/mod.rs \u2014 is it single-threaded?',
        'Run concurrent connection test: 10, 100, 1000 connections',
        'Compare Neqo server features against Cloudflare quiche and Google QUICHE',
        'Can Neqo work behind or alongside nginx?',
        'Does Neqo\u2019s limited server matter for our research, or is it sufficient for protocol testing?',
    ]:
        add_bullet(doc, q)

    # ── 4b ──
    doc.add_page_break()
    add_separator(doc)
    add_h2(doc, '4b. nginx & Server Architecture Comparison')
    add_body(doc,
        'nginx is the industry standard for load balancing. It supports QUIC/HTTP3 '
        '(since 1.25+) but does NOT implement preferred_address / server migration. '
        'Why not? What architecture is needed instead?')

    add_h3(doc, 'Architecture Comparison')
    add_body(doc, 'nginx (proxy in every packet):')
    add_body(doc, '    Client \u2194 nginx \u2194 Backend (nginx forwards ALL traffic, always in path)')
    add_body(doc, 'QUIC migration (proxy exits after handshake):')
    add_body(doc, '    Client \u2194 Front-end (handshake) \u2192 migration \u2192 Backend (direct)')

    add_table(doc, [
        ['Dimension', 'nginx Reverse Proxy', 'QUIC Migration'],
        ['Proxy in data path?', 'Yes, always', 'No \u2014 only handshake'],
        ['Per-packet overhead', 'Every packet forwarded', 'Zero after migration'],
        ['Server requirements', 'Only nginx needs QUIC', 'Both need QUIC stack'],
        ['Failure handling', 'Retry to another backend', 'One-shot migration'],
        ['Scalability ceiling', 'nginx is bottleneck', 'No bottleneck after'],
        ['Health checks', 'Built-in', 'Not built-in'],
        ['Routing logic', 'URL, headers, weight', 'Server decision at\nhandshake only'],
        ['Production ready?', 'Decades of use', 'Research (ours)'],
    ])

    add_h3(doc, 'Why nginx Doesn\u2019t Support Migration')
    add_body(doc,
        'nginx\u2019s architecture is fundamentally a proxy \u2014 it sits in the middle and '
        'forwards packets. Server-side migration requires the proxy to exit the data '
        'path, which contradicts nginx\u2019s core design. Additionally, preferred_address requires '
        'sharing TLS secrets with the backend, which nginx\u2019s security model does not allow. '
        'The architecture needed for migration is fundamentally different from proxy.')

    add_h3(doc, 'What Server Architecture IS Needed?')
    for p in [
        'A front-end willing to give up its connection (not a persistent proxy)',
        'Backend servers with full QUIC+TLS capability (not just TCP backends)',
        'A secure state transfer channel between front-end and backend',
        'Coordination mechanism for load decisions at handshake time',
        'This is a fundamentally different architecture than traditional reverse proxy',
    ]:
        add_bullet(doc, p)

    add_h3(doc, 'What to Check Further')
    for q in [
        'Can QUIC migration work alongside nginx? (nginx as primary, QUIC backend as preferred)',
        'Is a hybrid model practical: nginx for routing + QUIC migration for handoff?',
        'At what traffic volume does proxy bottleneck matter vs. migration?',
        'What operational features (health checks, retries) does migration need to replicate?',
        'Would CDN providers benefit from migration as alternative to proxy?',
    ]:
        add_bullet(doc, q)

    # ══════════════════════════════════════════════════════════════════
    # TASK 5: Attack Surface
    # ══════════════════════════════════════════════════════════════════
    doc.add_page_break()
    add_h1(doc, 'Task 5: Attack Surface vs. Benefits \u2014 Tradeoff Analysis')
    add_goal_box(doc,
        'Precisely define the attack surface opened by QUIC server-side migration and analyze '
        'the tradeoff: engineering advantages vs. security risks. Server-side migration is '
        'largely unimplemented specifically due to these concerns.',
        shade='FFF5F5')

    add_h2(doc, 'Why Server-Side Migration Is Rarely Implemented')
    add_body(doc,
        'Despite RFC 9000 Section 9.6, preferred_address is almost never implemented in '
        'production. Major implementations (Google QUICHE, Cloudflare quiche, nginx, msquic) '
        'all skip it. It opens an attack surface that didn\u2019t exist before.')

    add_h2(doc, 'Attack Surface Opened')
    add_bullet(doc,
        'Malicious server redirects client to attacker-controlled server. Client trusts '
        'preferred_address from TLS handshake \u2014 no independent verify.',
        bold_prefix='Connection Hijacking: ')
    add_bullet(doc,
        'Preferred serves different content. QUIC-Exfil paper: ML classifiers achieve only '
        '0-47% F1-score at detecting this.',
        bold_prefix='Covert Exfiltration: ')
    add_bullet(doc,
        '445 bytes of TLS secrets on the wire between servers. If intercepted, attacker '
        'decrypts all connection traffic.',
        bold_prefix='State Transfer Interception: ')
    add_bullet(doc,
        'Every server receiving migration state = compromise point.',
        bold_prefix='Trust Boundary Expansion: ')
    add_bullet(doc,
        'Client can choose not to use preferred_address (RFC 9000 \u00a79.6: \u2018MAY\u2019), '
        'but if it does, there is no mechanism to verify the preferred address belongs to the '
        'same legitimate operator. PATH_CHALLENGE validates reachability, not identity.',
        bold_prefix='No Migration Attestation: ')
    add_bullet(doc,
        'After migration, preferred controls what client receives.',
        bold_prefix='Content Substitution: ')

    add_h2(doc, 'Benefits Gained')
    add_bullet(doc,
        'Front-end exits data path after migration. Direct client-server communication.',
        bold_prefix='No Proxy Bottleneck: ')
    add_bullet(doc,
        'Faster than TCP reconnect + TLS 1.3 re-handshake (2 RTTs).',
        bold_prefix='Sub-Second Failover: ')
    add_bullet(doc,
        'Drain connections gracefully before shutdown.',
        bold_prefix='Zero-Downtime Maintenance: ')
    add_bullet(doc,
        'Lightweight TLS terminator + heavyweight app server.',
        bold_prefix='Vertical Separation: ')
    add_bullet(doc,
        'Mid-session migration to better server.',
        bold_prefix='Geographic Optimization: ')

    add_h2(doc, 'Risk/Benefit by Deployment Scenario')
    add_table(doc, [
        ['Scenario', 'Risk', 'Benefit', 'Verdict'],
        ['Same-subnet replicas\n(stateless LB)', 'LOW', 'HIGH \u2014 no proxy\nbottleneck', 'Favorable'],
        ['Vertical separation\n(same data center)', 'MEDIUM', 'HIGH \u2014 efficient\nresource usage', 'Likely favorable'],
        ['Cross-network\n(WAN migration)', 'HIGH', 'MEDIUM \u2014 geo\noptimization', 'Needs safeguards'],
        ['Cross-jurisdiction', 'VERY HIGH', 'LOW', 'Not recommended\nwithout attestation'],
    ])

    add_h2(doc, 'Research Questions')
    for i, q in enumerate([
        'Minimum attacker capability to exploit server-side migration?',
        'Can attack surface be reduced without losing engineering benefits?',
        'Which applications have acceptable risk/benefit?',
        'Should RFC require client consent or migration attestation?',
        'Can we propose a migration_attestation transport parameter?',
    ], 1):
        add_rq(doc, i, q)

    # ══════════════════════════════════════════════════════════════════
    # HOW TASKS CONNECT
    # ══════════════════════════════════════════════════════════════════
    doc.add_page_break()
    add_h1(doc, 'How the 5 Tasks Connect')
    add_bullet(doc,
        'Benchmark 4 state transfer backends in Docker. '
        'Produces concrete latency and reliability numbers.',
        bold_prefix='Task 1 (Docker + Backends): ')
    add_bullet(doc,
        'Investigate whether the preferred server can derive keys during the handshake '
        'instead of after. Document the ECDHE limitation.',
        bold_prefix='Task 2 (QUIC Proxy): ')
    add_bullet(doc,
        'Define use cases, check which server types work, analyze why stream-level LB '
        'is not feasible, look at multi-tier scenarios.',
        bold_prefix='Task 3 (Applications): ')
    add_bullet(doc,
        'Check what Neqo\u2019s server actually supports. Compare migration architecture '
        'against nginx reverse proxy.',
        bold_prefix='Task 4 (Neqo + Architecture): ')
    add_bullet(doc,
        'Document what attacks preferred_address enables and under what deployment '
        'conditions the risk is acceptable.',
        bold_prefix='Task 5 (Attack Surface): ')

    doc.save(OUTPUT_DOCX)
    print(f"Generated: {OUTPUT_DOCX}")


if __name__ == '__main__':
    build_doc()
